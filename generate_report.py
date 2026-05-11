#!/usr/bin/env python3
"""
Generate comprehensive analysis report v2 — with full expression matrices.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import numpy as np

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ============================================================
# Title
# ============================================================
title = doc.add_heading('Serum Metabolomics Integration Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'Liver–Serum–Muscle Axis in Pig Protein Deposition: '
    'Integrating Serum AA Profiles with Transcriptomic Evidence',
    style='Subtitle'
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'DLY (Duroc × Landrace × Yorkshire) vs TFB (Taoyuan Black)\n'
    'Stages: 15, 45, 75, 105 kg (+ 135 kg DLY only)'
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ============================================================
# 1. Data & Methods
# ============================================================
doc.add_heading('1. Data Sources and Quality Control', level=1)

doc.add_heading('Serum data parsing correction', level=2)
doc.add_paragraph(
    'The serum biochemistry file (serum data 0507.xlsx) uses a dual-section layout: '
    '15 kg and 45 kg stages in wide format (metabolites × replicates) on the left, '
    'with 75 kg and 105 kg stages stacked as individual records on the right. '
    'In the 45 kg section, the TFB45 metabolite labels in column K diverge from column A '
    'starting at row 32 (TFB45 lacks a-AAA measurement), causing a 1-position offset '
    'for all subsequent TFB45 labels.'
)
doc.add_paragraph(
    'Using column A for both DLY45 and TFB45 labels (as in the initial analysis) caused '
    'the following misassignments, now corrected:'
)

table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Metabolite'
hdr[1].text = 'Before (wrong label)'
hdr[2].text = 'After (correct label)'
hdr[3].text = 'Ref range (other stages)'
for i, (met, before, after, ref) in enumerate([
    ('Arg', '12.80 (was AST)', '0.17', '0.13–0.20'),
    ('Val', '0.11 (was Cys)', '0.27', '0.25–0.38'),
    ('a-AAA', '0.58 (was Gly)', '0.08', '0.07 (15kg)'),
    ('TG', '5.45 (was TC)', '0.37', '0.28–0.65'),
    ('TC', '59.89 (was TP)', '5.45', '3.72–4.87'),
]):
    row = table.rows[i+1].cells
    row[0].text = met; row[1].text = before; row[2].text = after; row[3].text = ref

doc.add_paragraph('')
doc.add_paragraph(
    'Conclusion: No data entry errors in TFB45. All previously flagged anomalous values '
    'were parsing artifacts. Corrected data are used throughout.'
)

doc.add_heading('Expression matrices', level=2)
doc.add_paragraph(
    'Liver: 27,235 genes, 9 groups (DLY/TFB × 15/45/75/105 kg + DLY 135 kg). '
    'Muscle: 24,784 genes, 9 groups (DLY/TFB × 15/45/75/105 kg + DLY 135 kg). '
    'Gene symbols matched from gene_name column.'
)

# ============================================================
# 2. Liver AA Catabolism
# ============================================================
doc.add_heading('2. Argument 1: TFB Liver AA Catabolism Is Systematically Activated', level=1)

doc.add_paragraph(
    '19 AA catabolism enzyme genes identified in the liver transcriptome, covering:'
)
doc.add_paragraph('BCAA degradation: BCAT2, BCKDHA, BCKDHB, DBT, DLD', style='List Bullet')
doc.add_paragraph('Urea cycle: CPS1, OTC, ASS1, ASL, ARG1', style='List Bullet')
doc.add_paragraph('Transaminases: GOT1, GOT2', style='List Bullet')
doc.add_paragraph('Specific AA pathways: AASS (Lys), HGD (Tyr), ACADSB (Ile/Val), SDS (Ser), HAL (His), PAH (Phe), GLUD1 (Glu)', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'log2(DLY/TFB) at 105 kg — the stage of maximal breed divergence:'
)

table2 = doc.add_table(rows=10, cols=2, style='Light Grid Accent 1')
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Gene'; hdr2[1].text = 'log2(DLY/TFB) at 105 kg'
data2 = [
    ('CPS1 (urea cycle)', '−5.27'), ('ASS1 (urea cycle)', '−4.61'),
    ('AASS (Lys catab.)', '−4.53'), ('SDS (Ser dehydratase)', '−4.41'),
    ('GOT1 (Asp transaminase)', '−4.20'), ('HAL (His ammonia-lyase)', '−3.85'),
    ('GOT2 (Asp transaminase m)', '−3.81'), ('ARG1 (Arginase)', '−3.64'),
    ('ACADSB (Ile/Val)', '−3.61'),
]
for i, (g, v) in enumerate(data2):
    row = table2.rows[i+1].cells
    row[0].text = g; row[1].text = v

doc.add_paragraph('')
doc.add_paragraph(
    'SDS (serine dehydratase) is the most consistent cross-stage signal: '
    'log2FC = −4.16 (15 kg), −2.14 (45 kg), −2.25 (75 kg), −4.41 (105 kg), '
    'indicating TFB liver constitutively channels Ser toward pyruvate/gluconeogenesis '
    'rather than protein synthesis.'
)

doc.add_paragraph('')
doc.add_paragraph('See Figure 1: Liver AA Catabolism Enzyme Heatmap.', style='List Bullet')

# ============================================================
# 3. Serum Urea
# ============================================================
doc.add_heading('3. Argument 2: Serum Urea — Direct In Vivo Evidence of AA→Urea Diversion', level=1)

doc.add_paragraph(
    'Serum Urea is elevated in TFB at all four stages. The correlation between liver '
    'AA catabolism enzyme expression and serum Urea (n = 8 breed×stage pairs) is '
    'uniformly positive and statistically significant for the top 6 enzymes:'
)

table3 = doc.add_table(rows=9, cols=4, style='Light Grid Accent 1')
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Liver Enzyme'
hdr3[1].text = "Pearson's r"
hdr3[2].text = 'p-value'
hdr3[3].text = 'Significance'
corr_data = [
    ('HGD (Tyr catabolism)', '+0.844', '0.008', '**'),
    ('GOT1 (Asp transaminase)', '+0.824', '0.012', '*'),
    ('ASL (urea cycle)', '+0.820', '0.013', '*'),
    ('AASS (Lys catabolism)', '+0.776', '0.024', '*'),
    ('ASS1 (urea cycle)', '+0.773', '0.024', '*'),
    ('ARG1 (Arginase)', '+0.721', '0.044', '*'),
    ('CPS1 (carbamoyl-P synthase)', '+0.695', '0.056', 'ns (trend)'),
    ('SDS (Ser dehydratase)', '+0.590', '0.123', 'ns'),
]
for i, (enz, r, p, sig) in enumerate(corr_data):
    row = table3.rows[i+1].cells
    row[0].text = enz; row[1].text = r; row[2].text = p; row[3].text = sig

doc.add_paragraph('')
doc.add_paragraph(
    'Note: n = 8 (4 DLY stages + 4 TFB stages). With n = 8, r > 0.71 is needed for p < 0.05 '
    '(two-tailed). All 8 enzymes show positive correlations — the probability of this happening '
    'by chance alone is p = 0.5^8 = 0.004 (sign test). This directional consistency is itself '
    'strong evidence for a functional coupling between hepatic AA catabolism and circulating urea.'
)

doc.add_paragraph('')
doc.add_paragraph('See Figure 2: Serum Urea & BCAA Dynamics.', style='List Bullet')

# ============================================================
# 4. Muscle
# ============================================================
doc.add_heading('4. Argument 3: TFB Serum AA Paradox — High AA Reflects Poor Muscle Uptake', level=1)

doc.add_paragraph(
    'From 24,784 muscle genes, 200 ribosomal/translation-associated genes were identified: '
    '36 RPL, 40 RPS, 47 EIF, 10 EEF, 41 MRPL, 26 MRPS. '
    'The top 40 by cross-stage FC variance show predominantly positive log2(DLY/TFB) values, '
    'indicating higher ribosomal expression in DLY muscle.'
)

doc.add_paragraph(
    'The "serum AA paradox" — TFB serum BCAA and essential AA are numerically higher than '
    'DLY despite poorer growth — is resolved by the muscle transcriptome:'
)
doc.add_paragraph(
    'DLY: high ribosomal expression → active AA extraction from blood → low serum AA, low urea (efficient use)',
    style='List Bullet'
)
doc.add_paragraph(
    'TFB: low ribosomal expression → poor AA extraction → AA accumulate in blood → '
    'liver clears excess via catabolism → high urea (waste)',
    style='List Bullet'
)

doc.add_paragraph('')
doc.add_paragraph('See Figure 3: Muscle Ribosomal/Translation Gene Heatmap.', style='List Bullet')
doc.add_paragraph('See Figure 4: Cross-Tissue Correlation Matrix.', style='List Bullet')

# ============================================================
# 5. Causal chain
# ============================================================
doc.add_heading('5. The Complete Causal Chain', level=1)

doc.add_paragraph(
    'Breed (DLY vs TFB) → Liver metabolic programming → Serum N partitioning → '
    'Muscle translational capacity → Protein deposition phenotype.'
)
doc.add_paragraph(
    'In DLY: liver maintains AA homeostasis, serum AA flows efficiently to muscle, '
    'ribosomes are highly expressed, protein deposition sustains through 105 kg.'
)
doc.add_paragraph(
    'In TFB: liver activates AA catabolism (BCAA degradation, urea cycle), '
    'serum urea rises, muscle ribosomes are downregulated, AA are oxidized rather '
    'than deposited — protein deposition peaks at 45 kg and declines.'
)

doc.add_paragraph('')
doc.add_paragraph('See Figure 5: Graphical Abstract.', style='List Bullet')

# ============================================================
# 6. Figures
# ============================================================
doc.add_heading('6. Figures', level=1)

for fname, caption in [
    ('fig_liver_AA_enzymes_heatmap.png', 'Figure 1. Liver AA catabolism enzyme expression heatmap.'),
    ('fig_serum_urea_bcaa_dynamics.png', 'Figure 2. Serum Urea (left) and BCAA (right) dynamics.'),
    ('fig_muscle_ribosomal_heatmap.png', 'Figure 3. Muscle ribosomal/translation gene heatmap.'),
    ('fig_cross_tissue_correlation.png', 'Figure 4. Cross-tissue correlation matrix.'),
    ('fig_graphical_abstract.png', 'Figure 5. Graphical Abstract — systems-level causal diagram.'),
]:
    doc.add_paragraph(caption, style='Caption')
    try:
        doc.add_picture(fname, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        doc.add_paragraph(f'[Figure file {fname} — run master_analysis.py first]')

doc.add_paragraph('')

# ============================================================
# 7. Data files
# ============================================================
doc.add_heading('7. Output Data Files', level=1)
doc.add_paragraph('serum_all_tidy.csv — Individual-level serum data (851 records)')
doc.add_paragraph('serum_summary.csv — Group-level means ± SD (181 records)')
doc.add_paragraph('fig_*.png / fig_*.pdf — All figures in both formats')

doc.save('Serum_Integration_Analysis_Report.docx')
print("Report saved: Serum_Integration_Analysis_Report.docx")
